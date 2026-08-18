"""DOCX parsing via python-docx: paragraphs + tables with section anchors."""

from __future__ import annotations

from io import BytesIO

import docx

from agent_server.knowledge_graph.schema import Chunk

_MIN_PARAGRAPH_CHARS = 20
_MIN_ROW_CHARS = 10


def parse_docx(file_content: bytes, document_id: str) -> list[Chunk]:
    """Parse a DOCX into one `Chunk` per non-trivial paragraph and one per
    table row, each carrying a `'para.<n>'` or `'table.<t> row.<r>'` locator."""
    document = docx.Document(BytesIO(file_content))
    chunks: list[Chunk] = []

    current_section: str | None = None
    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        if _is_heading(paragraph):
            current_section = text
        if len(text) < _MIN_PARAGRAPH_CHARS:
            continue
        chunks.append(
            Chunk(
                document_id=document_id,
                locator=f"para.{paragraph_index}",
                text=text,
                metadata={"paragraph": paragraph_index, "section": current_section},
            )
        )

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)
            if len(row_text.strip(" |")) < _MIN_ROW_CHARS:
                continue
            chunks.append(
                Chunk(
                    document_id=document_id,
                    locator=f"table.{table_index} row.{row_index}",
                    text=row_text,
                    metadata={"table": table_index, "row": row_index},
                )
            )
    return chunks


def _is_heading(paragraph) -> bool:
    style_name = getattr(paragraph.style, "name", "") or ""
    return style_name.lower().startswith("heading")
